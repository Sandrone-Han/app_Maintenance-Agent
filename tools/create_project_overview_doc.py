from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/设备维护排班系统项目概述与推进方案.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return paragraph


def add_status_table(doc):
    rows = [
        ("前端页面", "已完成主体页面，并开始切换为后端 API 数据源", "继续统一数据来源，减少 localStorage 依赖"),
        ("后端架构", "已搭建 NestJS + TypeScript + Oracle 连接基础", "补齐业务规则、统一异常处理和接口校验"),
        ("数据库", "已通过 Docker 搭建本地 Oracle，并创建核心业务表", "整理迁移、种子数据、真实数据库切换流程"),
        ("Excel 导入", "已支持人员角色技能表上传并写入 Oracle", "增加导入预览、字段校验、错误行提示"),
        ("排班能力", "已有排班任务、结果、日志相关数据结构", "实现规则版自动排班，再接入 AI 排班引擎"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1.35, 2.45, 2.7])
    headers = ["模块", "当前进度", "下一步重点"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="0B2545")
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_api_table(doc):
    rows = [
        ("健康检查", "GET /api/health", "确认后端服务是否启动"),
        ("数据库检查", "GET /api/db-check", "确认 Node 后端是否连接 Oracle"),
        ("班组人员", "/api/team-members", "人员 CRUD 与 Excel 导入"),
        ("班次类型", "/api/shift-types", "班次基础配置 CRUD"),
        ("出勤记录", "/api/attendance", "人员出勤状态 CRUD"),
        ("排班记录", "/api/team-schedule-records", "班组当前轮换状态查询"),
        ("排班任务", "/api/schedule-jobs", "创建排班任务并记录日志"),
        ("排班结果", "/api/schedule-results", "查询与导出排班结果"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1.45, 2.15, 2.9])
    for idx, header in enumerate(["能力", "接口范围", "用途"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="0B2545")
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_roadmap_table(doc):
    rows = [
        ("1", "稳定数据导入", "完善 Excel 校验、导入预览、错误行提示", "高"),
        ("2", "统一前后端数据源", "所有页面改为从 Oracle 读取和写入", "高"),
        ("3", "打通排班闭环", "配置参数 -> 创建任务 -> 生成结果 -> 前端查看", "高"),
        ("4", "实现规则版排班", "先用确定性规则实现上4休2、技能匹配、借调", "高"),
        ("5", "排班日志与任务状态", "后端记录节点日志，前端展示真实进度", "中"),
        ("6", "增强导出", "支持按筛选条件导出 CSV / Excel", "中"),
        ("7", "接入 AI 排班", "在规则版稳定后接入 AI 引擎做增强求解", "中"),
        ("8", "真实数据库切换", "通过环境变量接入生产或测试 Oracle", "中"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [0.55, 1.65, 3.35, 0.95])
    for idx, header in enumerate(["阶段", "推进方向", "具体工作", "优先级"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="0B2545")
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("设备维护排班系统项目概述与推进方案")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("基于当前前端、后端、Oracle 本地数据库与 Excel 导入功能的阶段性整理")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    add_heading(doc, "一、项目概述", 1)
    add_body(
        doc,
        "本项目是一个面向设备维护排班管理员和班组长的企业级中后台系统。系统目标是将班组人员、班次、出勤、班组轮换记录和排班结果从人工维护逐步升级为前后端分离、可连接真实数据库、后续可接入智能排班能力的业务系统。",
    )
    add_body(
        doc,
        "当前项目已经从纯前端 mock 原型推进到 React/Vite 前端、NestJS/Node.js 后端、Oracle 数据库和 Docker 本地测试环境组成的完整开发架构。",
    )

    add_heading(doc, "二、当前总体架构", 1)
    add_body(doc, "当前系统采用前后端分离架构，前端通过 HTTP API 调用后端，后端通过 Oracle 驱动访问本地 Docker Oracle 数据库。")
    for item in [
        "前端：React、Vite、TypeScript，负责页面展示、表单交互、Excel 上传入口和排班流程操作。",
        "后端：NestJS、TypeScript、oracledb，负责 API、数据库连接、业务数据维护和导入处理。",
        "数据库：Oracle Free，本地通过 Docker Compose 启动，用于模拟后期真实 Oracle 环境。",
        "数据流：前端页面 -> 后端 API -> Oracle 数据表 -> 返回页面展示或导出。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、目前已完成的内容", 1)
    add_status_table(doc)

    add_heading(doc, "四、后端当前能力", 1)
    add_body(doc, "后端已经完成第一阶段基础设施搭建，重点不是一次性完成全部业务，而是先形成可运行、可连接数据库、可导入测试数据的后端底座。")
    add_api_table(doc)

    add_heading(doc, "五、数据库与本地测试环境", 1)
    add_body(doc, "本地数据库已使用 Docker Compose 管理 Oracle 容器，后端通过环境变量读取连接信息。后期切换真实数据库时，原则上不需要修改业务代码，只需要调整数据库地址、服务名、账号和密码。")
    for item in [
        "本地服务名：FREEPDB1。",
        "本地用户：maintenance。",
        "主要业务表：TEAM_MEMBER、TEAM_MEMBER_SKILL、SHIFT_TYPE、ATTENDANCE_RECORD、TEAM_SCHEDULE_RECORD、SCHEDULE_JOB、SCHEDULE_JOB_LOG、SCHEDULE_RESULT。",
        "初始化方式：使用后端提供的 migrate、seed、verify 脚本创建和检查数据。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "六、当前阶段判断", 1)
    add_body(
        doc,
        "当前项目可以判断为“后端基础架构完成，核心数据接口初步完成，数据库本地测试环境完成，Excel 导入能力初步完成”。它已经不再是单纯前端原型，但还没有进入完整业务系统阶段。",
    )
    add_body(
        doc,
        "下一阶段的重点应放在数据稳定性、前后端联动一致性和排班业务闭环上。只有先把数据底座打稳，后续 AI 排班能力才容易接入和验证。",
    )

    add_heading(doc, "七、后续推进方向", 1)
    add_roadmap_table(doc)

    add_heading(doc, "八、建议推进顺序", 1)
    for item in [
        "完善 Excel 导入校验，确保人员数据能够稳定进入 Oracle。",
        "将前端所有页面统一切换到后端 API，避免 mock 与数据库数据并存导致混乱。",
        "实现排班任务的真实状态流转和日志记录。",
        "先实现规则版自动排班，覆盖班组轮换、技能匹配、出勤限制和借调规则。",
        "将排班结果保存到数据库，并在排班结果页查询和导出。",
        "规则版稳定后，再接入 AI 排班引擎作为增强能力。",
        "整理部署与真实数据库接入说明，形成后期上线基础。",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "九、近期最优先事项", 1)
    add_body(doc, "建议近期优先完成以下三件事：")
    for item in [
        "把 Excel 导入做稳定，包括字段校验、错误提示和重复数据处理。",
        "确认前端导入按钮、人员列表刷新、后端写库三者完整闭环。",
        "把人员、班次、出勤和班组排班记录统一切换为 Oracle 数据源。",
    ]:
        add_bullet(doc, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("设备维护排班系统项目文档")
    footer_run.font.name = "Microsoft YaHei"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor.from_string("666666")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
